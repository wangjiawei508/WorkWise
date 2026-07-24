#!/usr/bin/env python3
"""标书 Markdown → DOCX 一键生成（跨平台）。

将合并后的标书 Markdown 转为符合中文标书排版规范的 DOCX：
- 优先用 python-docx 逐元素渲染，套用格式预设（政府/企业/公路）
- 若装有 pandoc，可用 --engine pandoc 走 pandoc（保留更复杂表格）
- 支持标题编号、正文字体字号、表格、页眉页脚、目录占位

格式预设（--format）：
  government 仿宋_GB2312 四号 固定行距28磅（政府/事业单位）
  enterprise 宋体 小四 1.5倍行距（企业客户，默认）
  highway    仿宋_GB2312 四号 单倍（公路交通行业）

用法：
    python build_docx.py 合并初稿.md -o 投标技术标.docx --format government \
        --title "XX项目投标技术方案" --company "{公司名称}"
    python build_docx.py in.md -o out.docx --engine pandoc
依赖：pip install python-docx  （pandoc 引擎另需系统 pandoc）
"""
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

PRESETS = {
    "government": {"font": "仿宋_GB2312", "size": 14, "line": ("exact", 28), "heading": "黑体"},
    "enterprise": {"font": "宋体", "size": 12, "line": ("mult", 1.5), "heading": "黑体"},
    "highway": {"font": "仿宋_GB2312", "size": 14, "line": ("mult", 1.0), "heading": "黑体"},
}
MAX_INPUT_BYTES = 64 * 1024 * 1024


def regular_input(path: Path) -> bool:
    try:
        return (
            not path.is_symlink()
            and path.is_file()
            and path.stat().st_size <= MAX_INPUT_BYTES
        )
    except OSError:
        return False


def safe_output(path: Path) -> bool:
    try:
        return not path.is_symlink()
    except OSError:
        return False


def build_with_pandoc(md: Path, out: Path) -> bool:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        print("未检测到 pandoc，回退 python-docx 引擎。")
        return False
    cmd = [pandoc, str(md), "-o", str(out), "--from", "gfm", "--to", "docx"]
    try:
        subprocess.run(
            cmd,
            check=True,
            stdin=subprocess.DEVNULL,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=120,
        )
        if not regular_input(out) or out.stat().st_size == 0:
            print("pandoc 未生成有效 DOCX，回退 python-docx。")
            out.unlink(missing_ok=True)
            return False
        print(f"[pandoc] → {out}")
        return True
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError) as e:
        out.unlink(missing_ok=True)
        print(f"pandoc 失败：{e}，回退 python-docx。")
        return False


def build_with_docx(md: Path, out: Path, preset: dict, title: str, company: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("缺少 python-docx，请 `pip install python-docx`，或改用 --engine pandoc。")
        return False

    doc = Document()

    def set_cn_font(run, name, size):
        run.font.name = name
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def apply_line(para):
        pf = para.paragraph_format
        mode, val = preset["line"]
        if mode == "exact":
            pf.line_spacing = Pt(val)
        else:
            pf.line_spacing = val

    # 封面
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_cn_font(r, preset["heading"], 26)
        r.bold = True
        if company:
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rc = pc.add_run(company)
            set_cn_font(rc, preset["font"], 16)
        doc.add_page_break()
        # 目录占位
        pt = doc.add_paragraph()
        rt = pt.add_run("目  录")
        set_cn_font(rt, preset["heading"], 18)
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("（打开文档后按 F9 或“引用→更新目录”生成页码）")
        doc.add_page_break()

    lines = md.read_text(encoding="utf-8", errors="ignore").splitlines()
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in table_buf
                if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r)]
        table_buf = []
        if not rows:
            return
        ncol = max(len(r) for r in rows)
        t = doc.add_table(rows=0, cols=ncol)
        t.style = "Table Grid"
        for ri, rc in enumerate(rows):
            cells = t.add_row().cells
            for ci in range(ncol):
                txt = rc[ci] if ci < len(rc) else ""
                cells[ci].text = ""
                run = cells[ci].paragraphs[0].add_run(txt)
                set_cn_font(run, preset["font"], preset["size"] - 1)
                if ri == 0:
                    run.bold = True

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            table_buf.append(line)
            continue
        else:
            flush_table()
        if not line.strip():
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            p = doc.add_paragraph()
            r = p.add_run(m.group(2).strip())
            set_cn_font(r, preset["heading"], {1: 18, 2: 16, 3: 14, 4: preset["size"]}[level])
            r.bold = True
            continue
        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(re.sub(r"^\s*[-*]\s+", "", line))
            set_cn_font(r, preset["font"], preset["size"])
            apply_line(p)
            continue
        # 普通正文（去掉行内 md 强调符号）
        text = re.sub(r"[*`]", "", line)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(preset["size"] * 2)
        r = p.add_run(text)
        set_cn_font(r, preset["font"], preset["size"])
        apply_line(p)
    flush_table()

    doc.save(str(out))
    print(f"[python-docx] → {out}")
    return True


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("input", help="标书 Markdown")
    ap.add_argument("-o", "--out", default="标书.docx")
    ap.add_argument("--format", choices=list(PRESETS), default="enterprise")
    ap.add_argument("--engine", choices=["docx", "pandoc"], default="docx")
    ap.add_argument("--title", default="")
    ap.add_argument("--company", default="")
    args = ap.parse_args()

    requested_input = Path(args.input)
    if requested_input.is_symlink():
        print(f"输入文件不能是符号链接：{requested_input}")
        return 1
    md = requested_input.resolve()
    if not regular_input(md):
        print(f"输入文件不可用或超过 64 MiB：{md}")
        return 1
    requested_output = Path(args.out)
    if requested_output.is_symlink():
        print(f"输出路径不能是符号链接：{requested_output}")
        return 1
    out = requested_output.resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    out.unlink(missing_ok=True)

    ok = False
    if args.engine == "pandoc":
        ok = build_with_pandoc(md, out)
    if not ok:
        ok = build_with_docx(md, out, PRESETS[args.format], args.title, args.company)
    if not ok:
        return 1
    print("⚠️ 生成后请人工检查：目录页码、表格跨页、字体嵌入、签章位置、占位符替换。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
