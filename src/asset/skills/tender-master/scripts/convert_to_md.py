#!/usr/bin/env python3
"""把招标文件（PDF/DOCX/TXT）转换为 Markdown，便于 AI 阅读。

本脚本只是内置文档解析器不可用时的本地后备工具。转换失败会返回非零
退出码，且不会留下伪装成有效结果的 Markdown 文件。
用法：
    python convert_to_md.py 招标文件.pdf --output 招标文件.md
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


MAX_INPUT_BYTES = 64 * 1024 * 1024


class ConversionError(RuntimeError):
    pass


def convert_pdf(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        raise ConversionError(
            "缺少 pdfplumber，无法安全解析 PDF；请使用 WorkWise 内置文档解析器。"
        )
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            parts.append(f"\n<!-- 第{i}页 -->\n{text}")
            for table in page.extract_tables() or []:
                parts.append(_table_to_md(table))
    body = "\n".join(parts).strip()
    if not body:
        raise ConversionError(
            "PDF 未提取到文本，可能是扫描件；请启用高精度/OCR 文档引擎。"
        )
    return body


def convert_docx(path: Path) -> str:
    if path.suffix.lower() == ".doc":
        raise ConversionError(
            "旧版二进制 .doc 不能由 python-docx 安全解析；"
            "请使用 WorkWise 文档解析器或受信任的 Office 应用转换为 .docx。"
        )
    try:
        import docx  # type: ignore
    except ImportError:
        raise ConversionError(
            "缺少 python-docx，无法安全解析 DOCX；请使用 WorkWise 内置文档解析器。"
        )
    try:
        document = docx.Document(str(path))
    except Exception as error:
        raise ConversionError(f"DOCX 无法解析：{error}") from error
    lines = []
    for para in document.paragraphs:
        style = (para.style.name or "").lower()
        text = para.text.strip()
        if not text:
            continue
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        else:
            lines.append(text)
    for table in document.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        lines.append(_table_to_md(rows))
    body = "\n\n".join(lines).strip()
    if not body:
        raise ConversionError("DOCX 中没有可提取正文。")
    return body


def _table_to_md(rows) -> str:
    rows = [r for r in rows if r]
    if not rows:
        return ""
    header = rows[0]
    md = ["| " + " | ".join(str(c or "") for c in header) + " |",
          "| " + " | ".join("---" for _ in header) + " |"]
    for r in rows[1:]:
        md.append("| " + " | ".join(str(c or "") for c in r) + " |")
    return "\n".join(md)


def main() -> int:
    ap = argparse.ArgumentParser(description="招标文件转 Markdown")
    ap.add_argument("input", help="PDF/DOCX/DOC/TXT 文件路径")
    ap.add_argument("--output", help="输出 md 路径（默认同名 .md）")
    args = ap.parse_args()

    src = Path(args.input)
    if not src.exists():
        print(f"文件不存在：{src}")
        return 1
    if src.is_symlink() or not src.is_file():
        print(f"输入必须是普通文件且不能是符号链接：{src}")
        return 1
    if src.stat().st_size > MAX_INPUT_BYTES:
        print(f"输入文件超过 {MAX_INPUT_BYTES // (1024 * 1024)} MiB 安全上限：{src}")
        return 1
    suffix = src.suffix.lower()
    try:
        if suffix == ".pdf":
            body = convert_pdf(src)
        elif suffix in {".docx", ".doc"}:
            body = convert_docx(src)
        elif suffix in {".txt", ".md"}:
            body = src.read_text(encoding="utf-8", errors="strict").strip()
            if not body:
                raise ConversionError("输入文本为空。")
        else:
            print(f"暂不支持的格式：{suffix}（支持 pdf/doc/docx/txt/md）")
            return 1
    except (ConversionError, UnicodeError, OSError) as error:
        print(f"转换失败：{error}", file=sys.stderr)
        return 1

    out = Path(args.output) if args.output else src.with_suffix(".md")
    if out.exists() and out.is_symlink():
        print(f"输出不能是符号链接：{out}", file=sys.stderr)
        return 1
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(body, encoding="utf-8")
    print(f"已生成：{out}（{len(body)} 字符）")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
